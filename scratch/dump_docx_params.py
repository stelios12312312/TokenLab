import docx
import re

doc = docx.Document("docs_final/Z1_TOKEN_LIFECYCLE_V2.docx")
print(f"Total tables: {len(doc.tables)}")

for t_idx, table in enumerate(doc.tables):
    if not table.rows:
        continue
    headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
    
    # Find parameter column
    param_col_idx = -1
    for idx, h in enumerate(headers):
        if "parameter" in h:
            param_col_idx = idx
            break
            
    if param_col_idx == -1:
        continue
        
    val_col_idx = -1
    for idx, h in enumerate(headers):
        if h in ["value", "default", "baseline / default", "baseline / default value"]:
            val_col_idx = idx
            break
            
    if val_col_idx == -1:
        continue
        
    print(f"\nTable {t_idx} (headers: {headers}):")
    for row in table.rows[1:]:
        if len(row.cells) <= max(param_col_idx, val_col_idx):
            continue
        param_text = row.cells[param_col_idx].text.strip()
        val_text = row.cells[val_col_idx].text.strip()
        print(f"  - '{param_text}' : '{val_text}'")
