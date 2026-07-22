import sys
from docx import Document

def reset_document(doc):
    body = doc.element.body
    to_remove = []
    found = False
    
    for child in list(body):
        # check if it is a paragraph or a table or another element
        # paragraph tag ends with 'p', table tag ends with 'tbl'
        if child.tag.endswith('p'):
            text = "".join(node.text for node in child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
            if text.strip().startswith("14. Technical Appendix"):
                found = True
        
        if found:
            to_remove.append(child)
            
    if to_remove:
        print(f"Found {len(to_remove)} elements to remove.")
        for elem in to_remove:
            body.remove(elem)
        return True
    return False

if __name__ == "__main__":
    docx_path = "scratch/Z1_TOKEN_LIFECYCLE_V2_copy.docx"
    import shutil
    shutil.copy("docs_final/Z1_TOKEN_LIFECYCLE_V2.docx", docx_path)
    
    doc = Document(docx_path)
    print("Total paragraphs before reset:", len(doc.paragraphs))
    reset_document(doc)
    doc.save(docx_path)
    
    # Read back to verify
    doc2 = Document(docx_path)
    print("Total paragraphs after reset:", len(doc2.paragraphs))
    
    # Check if "14. Technical Appendix" is in the text
    found_14 = False
    for p in doc2.paragraphs:
        if "14. Technical Appendix" in p.text:
            found_14 = True
    print("Is Section 14 still present?", found_14)
