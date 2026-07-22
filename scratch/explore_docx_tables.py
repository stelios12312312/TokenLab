import docx
import sys

def explore(path):
    doc = docx.Document(path)
    print(f"Total tables: {len(doc.tables)}")
    for idx, table in enumerate(doc.tables):
        print(f"\n--- Table {idx} ---")
        rows = len(table.rows)
        cols = len(table.columns) if rows > 0 else 0
        print(f"Dimensions: {rows}x{cols}")
        
        # Print first row (header candidate)
        if rows > 0:
            header = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells]
            print(f"Header: {header}")
            
            # Print first 2 data rows
            for r_idx in range(1, min(rows, 3)):
                row_vals = [cell.text.strip().replace('\n', ' ') for cell in table.rows[r_idx].cells]
                print(f"  Row {r_idx}: {row_vals}")

if __name__ == "__main__":
    explore(sys.argv[1] if len(sys.argv) > 1 else "docs_final/Z1_TOKEN_LIFECYCLE_V2.docx")
