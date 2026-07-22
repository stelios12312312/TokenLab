import re
import sys
from docx import Document

def clean_text(text):
    # Remove markdown link syntax: [text](url) -> text
    # Supports up to one level of nested parentheses in url (e.g. (Personal))
    text = re.sub(r'\[([^\]]+)\]\((?:[^()]+|\([^()]*\))*\)', r'\1', text)
    # Remove bold/italic/code markup
    text = text.replace('**', '').replace('`', '').replace('*', '')
    return text

def add_bullet_list_item(doc, text):
    if 'List Bullet' in doc.styles:
        doc.add_paragraph(clean_text(text), style='List Bullet')
    else:
        doc.add_paragraph(f"• {clean_text(text)}")

def add_numbered_list_item(doc, text, num_str):
    if 'List Number' in doc.styles:
        doc.add_paragraph(clean_text(text), style='List Number')
    else:
        doc.add_paragraph(f"{num_str}. {clean_text(text)}")

def set_table_style(doc, table):
    try:
        table.style = 'Table Grid'
    except Exception as e:
        print(f"Warning: Could not set table style to 'Table Grid' ({e}). Finding alternative table style...")
        from docx.enum.style import WD_STYLE_TYPE
        applied = False
        for s in doc.styles:
            if s.type == WD_STYLE_TYPE.TABLE:
                try:
                    table.style = s.name
                    applied = True
                    break
                except Exception:
                    pass
        if not applied:
            print("Could not apply any custom table style. Leaving as default.")

def append_md_to_docx(md_path, docx_path):
    print(f"Reading markdown from: {md_path}")
    print(f"Opening DOCX file: {docx_path}")
    doc = Document(docx_path)
    
    # Add a page break to separate the new content cleanly
    doc.add_page_break()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_data = []
    
    for line in lines:
        line_str = line.strip()
        
        # Parse tables
        if line_str.startswith('|'):
            # Skip separator line
            if re.match(r'^\|[\s:-|]+$', line_str):
                continue
            # Parse row
            row_cells = [cell.strip() for cell in line_str.split('|')[1:-1]]
            table_data.append(row_cells)
            in_table = True
            continue
        else:
            if in_table:
                # Write collected table data to docx
                if table_data:
                    print(f"Writing table with {len(table_data)} rows and {len(table_data[0])} columns")
                    cols = len(table_data[0])
                    table = doc.add_table(rows=0, cols=cols)
                    set_table_style(doc, table)
                    for r_idx, row in enumerate(table_data):
                        row_cells_objs = table.add_row().cells
                        for c_idx, cell in enumerate(row):
                            row_cells_objs[c_idx].text = clean_text(cell)
                in_table = False
                table_data = []
                
        if not line_str:
            continue
            
        # Parse headings
        heading_match = re.match(r'^(#{1,6})\s+(.*)$', line_str)
        if heading_match:
            hashes, text = heading_match.groups()
            level = len(hashes)
            doc.add_heading(clean_text(text), level=min(level, 4))
            continue
            
        # Parse bullet lists
        list_match = re.match(r'^[-*]\s+(.*)$', line_str)
        if list_match:
            text = list_match.group(1)
            add_bullet_list_item(doc, text)
            continue
            
        # Parse numbered list
        num_list_match = re.match(r'^(\d+)\.\s+(.*)$', line_str)
        if num_list_match:
            num_str, text = num_list_match.groups()
            add_numbered_list_item(doc, text, num_str)
            continue
            
        # Regular paragraph
        doc.add_paragraph(clean_text(line_str))
        
    # Final check for tables if file ends on table
    if in_table and table_data:
        cols = len(table_data[0])
        table = doc.add_table(rows=0, cols=cols)
        set_table_style(doc, table)
        for row in table_data:
            row_cells_objs = table.add_row().cells
            for c_idx, cell in enumerate(row):
                row_cells_objs[c_idx].text = clean_text(cell)
                
    doc.save(docx_path)
    print("DOCX file saved successfully!")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python3 append_stelios_answers.py <md_path> <docx_path>")
        sys.exit(1)
    append_md_to_docx(sys.argv[1], sys.argv[2])
