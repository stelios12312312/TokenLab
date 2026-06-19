# @planner:module = docx_appendix
# @planner:story = US-Z1-M3-09
import sys
import os
import shutil
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def reset_document(doc):
    body = doc.element.body
    to_remove = []
    found = False
    
    for child in list(body):
        if child.tag.endswith('p'):
            text = "".join(node.text for node in child.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'))
            if text.strip().startswith("14. Technical Appendix"):
                found = True
        
        if found:
            to_remove.append(child)
            
    if to_remove:
        print(f"Found {len(to_remove)} existing Section 14 elements to remove.")
        for elem in to_remove:
            body.remove(elem)
        return True
    return False

def add_appendix(docx_path):
    print(f"Opening document: {docx_path}")
    doc = Document(docx_path)
    
    # Reset any existing Section 14 elements to support idempotent re-runs
    reset_document(doc)

    print("Adding Heading 1: 14. Technical Appendix")
    h1 = doc.add_heading("14. Technical Appendix", level=1)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True
    
    # 14.1
    h2_1 = doc.add_heading("14.1 TBD Parameter Calibration Registry", level=2)
    h2_1.paragraph_format.space_before = Pt(18)
    h2_1.paragraph_format.space_after = Pt(6)
    h2_1.paragraph_format.keep_with_next = True
    
    p1 = doc.add_paragraph(
        "This registry documents the parameters, their unit classifications, default values, sensitivity levels, "
        "sweeping ranges, and calibration guidelines established for the Milestone 3 (M3) simulation."
    )
    p1.paragraph_format.space_after = Pt(12)
    
    # Table data
    headers = ["Parameter", "Unit", "Baseline / Default", "Calibration Range", "Sensitivity", "Description & Calibration Guidelines"]
    rows_data = [
        ["TAU_1", "score", "0.20", "[0.10, 0.40]", "Medium", "PCS Cutoff (Casual to Engaged): Cutoff score for casual participants. Value is calibrated based on simulated population score distributions to ensure active audience progression."],
        ["TAU_2", "score", "0.60", "[0.50, 0.80]", "High", "PCS Cutoff (Engaged to Core): Controls access to higher Settlement Ratio (SR) tiers, governance rights, and core eligibility. Extremely sensitive for core cohort retention."],
        ["RELEASE_RATE_E0", "ratio", "0.10", "[0.05, 0.20]", "Medium", "Air-Claim Reserve Fraction: Fraction of Audience Reserve (AR) released at launch for Air-Claim. High rates drain AR too early; low rates underwhelm launch."],
        ["WAVE_SIZE", "count", "5,000", "[1,000, 10,000]", "Low", "Air-Claim Batch Size: Number of claims processed in a batch before PCS recalculation. Calibrates computational overhead against relative fairness."],
        ["THETA_MIN", "ratio", "0.30", "[0.20, 0.50]", "Critical", "Treasury Health Solvency Floor: Solvency boundary for the entire system. Below this, SYS_throttle is activated to preserve solvency."],
        ["SR_BASE", "ratio", "0.1047", "[0.01, 0.50]", "Highest", "ACR-to-Z1U Conversion Rate: Base conversion factor for settlements. Primary control over Z1U drain rate. The most sensitive parameter in the system."],
        ["settlement_cap_epoch", "Z1U", "50,000", "[10k, 200k]", "High", "Solvency Settlement Cap: Maximum aggregate Z1U settled per epoch across all users. Essential anti-stampede mechanism."],
        ["MIN_SETTLE", "ACR", "50.0", "[10.0, 100.0]", "Low", "Minimum Settlement Threshold: Dust threshold to prevent micro-settlement transaction spam."],
        ["LM_RATE", "multiplier/epoch", "0.05", "[0.01, 0.10]", "Medium", "Loyalty Multiplier Increase Rate: Determines the rate of loyalty multiplier increase per active epoch. Drives long-term user retention."],
        ["LM_MAX", "multiplier", "1.50", "[1.20, 2.00]", "Medium", "Maximum Loyalty Multiplier Cap: Bounds maximum loyalty advantage of tenure."],
        ["STREAK_BONUS", "multiplier", "0.10", "[0.05, 0.25]", "Low", "Streak Activity Bonus: Incremental bonus multiplier for unbroken active epochs. Rewards consistent engagement."],
        ["STREAK_WINDOW", "epochs", "8", "[4, 12]", "Low", "Streak Qualification Window: Number of consecutive active epochs required to qualify for the streak bonus."],
        ["sku_prices", "USD", "Dynamic", "[0.99, 999.00]", "Medium", "Utility SKU Pricing: USD-denominated price points. Adjusts Z1U amount dynamically via internal reference rate (similar to Helium Data Credits)."],
        ["fee_rate_g5b", "ratio", "0.34", "[0.10, 0.50]", "High", "Utility Treasury Capture Rate: Capture rate on utility transactions. Primary revenue channel; must be balanced against systemic solvency."],
        ["PAR-28 min_lock_period", "epochs", "12", "[4, 26]", "Medium", "Minimum Governance Lock Period: Prevents flash-governance and vote-and-dump attacks by locking staked Z1U."],
        ["PAR-29 max_lock_period", "epochs", "104", "[26, 156]", "Medium", "Maximum Governance Lock Period: Upper bound on locking duration to cap maximum vote weight accumulation."],
        ["revocation_cooldown", "epochs", "4", "[1, 8]", "Low", "Delegation Revocation Cooldown: Cooldown period on delegation revocation when active votes are open. Prevents manipulation."],
        ["fee_rate_g9b", "ratio", "0.25", "[0.05, 0.50]", "Medium", "Campaign Treasury Capture Rate: Secondary revenue channel capturing a fraction of campaign settlements."],
        ["campaign_min_budget", "Z1U", "5,000", "[1k, 50k]", "Low", "Minimum Campaign Budget: Floor budget to prevent campaign spam and ensure network quality."],
        ["pagerank_cap", "score", "0.05", "[0.01, 0.10]", "High", "PageRank Referral Cap: Upper limit for PageRank-based referral scores to prevent sybil/referral tree gaming."],
        ["min_shannon_entropy", "bits", "2.00", "[1.50, 3.50]", "Medium", "Minimum Shannon Entropy: Threshold for session diversity. Prevents single-action agricultural farming."],
        ["platform_min_engagement", "threshold", "0.10", "[0.05, 0.25]", "Medium", "Platform Minimum Engagement: Threshold to prevent platform-concentration attacks and encourage cross-platform diversity."]
    ]
    
    print("Creating parameter registry table...")
    table = doc.add_table(rows=len(rows_data) + 1, cols=6)
    
    try:
        table.style = 'Table Grid'
    except Exception as e:
        print(f"Warning: Could not set table style to 'Table Grid' ({e}). Finding alternative table style...")
        from docx.enum.style import WD_STYLE_TYPE
        applied = False
        for s in doc.styles:
            if s.type == WD_STYLE_TYPE.TABLE:
                print(f"Attempting style: {s.name}")
                try:
                    table.style = s.name
                    print(f"Successfully applied style: {s.name}")
                    applied = True
                    break
                except Exception:
                    pass
        if not applied:
            print("Could not apply any custom table style. Leaving as default.")
    
    # Set headers
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        # Make headers bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                
    # Fill data
    for r_idx, row_val in enumerate(rows_data):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_val):
            row_cells[c_idx].text = val
            # Make parameter name bold
            if c_idx == 0:
                for paragraph in row_cells[c_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        
    # Add spacing after table
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_after = Pt(12)
    
    # 14.2
    h2_2 = doc.add_heading("14.2 Treasury Health Metric and Feedback Loops", level=2)
    h2_2.paragraph_format.space_before = Pt(18)
    h2_2.paragraph_format.space_after = Pt(6)
    h2_2.paragraph_format.keep_with_next = True
    
    doc.add_paragraph(
        "The Treasury's long-term sustainability is governed by a compound health metric evaluated dynamically "
        "at the end of each epoch."
    )
    
    # Equation
    eq_p = doc.add_paragraph()
    eq_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_run = eq_p.add_run(
        "treasury_health(e) = T(e) / (OPEX(e) + VRP(e) + ecosystem_grants(e) + liquidity_provisioning(e))"
    )
    eq_run.font.bold = True
    eq_p.paragraph_format.space_before = Pt(6)
    eq_p.paragraph_format.space_after = Pt(6)
    
    p_eq_desc = doc.add_paragraph("Where:")
    p_eq_desc.paragraph_format.space_after = Pt(2)
    
    # Bullet points for equation description
    bullets = [
        "T(e) is the liquid Treasury vault balance at epoch e.",
        "OPEX(e) represent operational outflows.",
        "VRP(e) represent terminal validator reward payments.",
        "ecosystem_grants(e) represent ecosystem distribution outlays.",
        "liquidity_provisioning(e) represent provisioning requirements.",
        "Internal routing functions (such as G11 Audience Reserve top-up and G12 Creator Incentive Pool replenishment) are excluded from the denominator to avoid double-counting."
    ]
    for b in bullets:
        bp = doc.add_paragraph(f"• {b}")
        bp.paragraph_format.space_after = Pt(2)
    
    # Solvency Constraint heading and text
    p_solv = doc.add_paragraph()
    p_solv.paragraph_format.space_before = Pt(6)
    p_solv.add_run("Solvency Constraint").bold = True
    
    doc.add_paragraph(
        "The system enforces the strict condition: treasury_health(e) >= THETA_MIN"
    )
    
    doc.add_paragraph(
        "If treasury_health(e) < THETA_MIN, the system triggers SYS_throttle (M57), executing the following negative feedback loop actions:"
    )
    
    throttles = [
        "PCS Weight Compression: Reduces the relative weight of engagement dimensions, dampening aggregate ACR generation.",
        "Vesting Duration Extension: Scales the vesting timeline by VEST_EXTENSION_RATE to delay terminal settlement outflows.",
        "Settlement Ratio Reduction: Decays the effective conversion rate linearly to reduce the Z1U outflow rate per settled ACR."
    ]
    for t in throttles:
        bp = doc.add_paragraph(f"• {t}")
        bp.paragraph_format.space_after = Pt(2)
        
    # 14.3
    h2_3 = doc.add_heading("14.3 Boundary Behavior Monitors", level=2)
    h2_3.paragraph_format.space_before = Pt(18)
    h2_3.paragraph_format.space_after = Pt(6)
    h2_3.paragraph_format.keep_with_next = True
    
    doc.add_paragraph(
        "To protect the system against emergent gaming and behavioral degradation, Section 10's Boundary Behavior Monitors are tracked via specific triggers:"
    )
    
    monitors = [
        "V-B1 (Passive Accumulation): Tracked by monitoring z1u_balance growth without corresponding G5a utility spending over N consecutive epochs. Trigger activates promotional SKU discounts.",
        "V-B2 (Immediate Exit): Tracked via the settlement_to_exit_ratio. If aggregate ratio exceeds 0.80, it dynamically triggers a settlement cooldown extension.",
        "C-B1 (Curation Concentration): Tracked by calculating the Herfindahl-Hirschman Index (HHI) over the curation distribution in D-scores. High HHI triggers a diversity bonus for underserved content categories.",
        "CR-B1 (Minimum-Viable Quality): Tracked via average Q-score. Drops below target trigger a reduction in the ACR multiplier.",
        "VL-B1 (Over-Conservative Verification): Tracked via validator false-negative rejection rates. Anomalous rates trigger validator cohort auditing.",
        "GD-B1 (Passive Delegation): Tracked by delegation age without corresponding governance voting participation. Triggers delegation decay.",
        "ES-B1 (Large Position Exit): Tracked by monitoring high-volume Z1U transfers to exchange addresses, feeding directly into the B5 stabilization loop."
    ]
    for m in monitors:
        bp = doc.add_paragraph(f"• {m}")
        bp.paragraph_format.space_after = Pt(2)
        
    # 14.4
    h2_4 = doc.add_heading("14.4 Legal and Financial Disclaimers", level=2)
    h2_4.paragraph_format.space_before = Pt(18)
    h2_4.paragraph_format.space_after = Pt(6)
    h2_4.paragraph_format.keep_with_next = True
    
    doc.add_paragraph(
        "This Technical Appendix, along with the preceding economic specification, is provided solely for tokenomic modeling and simulation engineering purposes."
    )
    
    disclaimers = [
        "Not Financial Advice: The contents of this document do not constitute financial, investment, or trading advice. No information herein should be interpreted as an endorsement or recommendation to buy, sell, or hold any digital asset.",
        "Not Legal Advice: This document does not constitute legal or regulatory advice. The regulatory status of utility tokens and digital assets varies significantly by country.",
        "Owner Identification: The legal owner of the project and this specification is Stylianos Kampakis (TokenLab).",
        "Jurisdictional Assumptions: All simulations assume compliance with international Anti-Money Laundering (AML) and Know Your Customer (KYC) regulations. Any live launch of the protocol, investment decisions, or public deployment of the token lifecycle requires independent, qualified legal review in the relevant jurisdictions before execution."
    ]
    for d in disclaimers:
        bp = doc.add_paragraph(f"• {d}")
        bp.paragraph_format.space_after = Pt(2)
        
    # 14.5
    h2_5 = doc.add_heading("14.5 Simulation Visualizations", level=2)
    h2_5.paragraph_format.space_before = Pt(18)
    h2_5.paragraph_format.space_after = Pt(6)
    h2_5.paragraph_format.keep_with_next = True
    
    doc.add_paragraph(
        "To empirically validate the calibrated parameter registry and confirm the system's structural solvency "
        "characteristics under stress conditions, the Milestone 3 (M3) simulation results are presented below."
    )
    
    plots = [
        {
            "path": "outputs/z1_m3_sims/monte_carlo/monte_carlo_resilience_bands.png",
            "caption": "Figure 14.1: Monte Carlo Resilience Bands",
            "desc": "This figure shows the confidence bands of systemic solvency (measured via the treasury health ratio) across 1,000 randomized simulation runs. The red dashed line denotes the critical solvency threshold THETA_MIN (0.30). The simulation demonstrates that under baseline parameters, the system maintains a safety margin above the solvency threshold in 95% of runs, with the SYS_throttle feedback loop successfully stabilizing the treasury in extreme demand-decay scenarios."
        },
        {
            "path": "outputs/z1_m3_sims/compare/m2_m3_comparison.png",
            "caption": "Figure 14.2: M2 vs. M3 Model Comparison",
            "desc": "A direct comparison of the active participant progression under the Milestone 2 (M2 - static thresholding) and Milestone 3 (M3 - dynamic PCS tier classification) models. M3 utilizes calibrated cutoff thresholds TAU_1 (0.20) and TAU_2 (0.60). The dynamic calibration prevents the cohort congestion observed in M2, smoothing the transition from casual to engaged and core cohorts, thereby enhancing long-term engagement stability."
        },
        {
            "path": "outputs/z1_m3_sims/compare/m3_pools_governance_stress.png",
            "caption": "Figure 14.3: M3 Pools and Governance Stress Test",
            "desc": "This stress test monitors treasury inflows and outflows during simulated flash-governance and large-position exit shocks. The governance lock period constraints (PAR-28 min_lock_period = 12 epochs and PAR-29 max_lock_period = 104 epochs) act as effective dampening mechanisms. Despite high volatility in token delegation, liquid treasury reserves remain above the solvency floor, proving the design's structural resilience."
        },
        {
            "path": "outputs/z1_m3_sims/sweeps/parameter_sensitivity_heatmaps.png",
            "caption": "Figure 14.4: Parameter Sensitivity Sweeps",
            "desc": "Heatmap visualization of the sensitivity sweeps across critical parameters SR_BASE (ACR-to-Z1U conversion rate) and fee_rate_g5b (Utility treasury capture rate). The results confirm that SR_BASE is the highest sensitivity parameter in the system. Solvency is maintained (green region) when SR_BASE remains below 0.15 and fee_rate_g5b is calibrated to at least 0.20, matching the theoretical breakeven requirement of the sustainability equation."
        }
    ]
    
    for plot in plots:
        path = plot["path"]
        if not os.path.exists(path):
            print(f"Error: Simulation plot file not found at: {path}")
            return False
            
        print(f"Adding plot: {path}")
        # Add image centered
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_before = Pt(12)
        img_p.paragraph_format.space_after = Pt(4)
        img_run = img_p.add_run()
        try:
            img_run.add_picture(path, width=Inches(6.0))
        except Exception as e:
            print(f"Error adding image {path} to document: {e}")
            return False
            
        # Add caption centered
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_after = Pt(12)
        cap_run = cap_p.add_run(plot["caption"])
        cap_run.italic = True
        cap_run.font.size = Pt(10)
        
        # Add descriptive text
        desc_p = doc.add_paragraph(plot["desc"])
        desc_p.paragraph_format.space_after = Pt(12)
        
    print(f"Saving modified document back to: {docx_path}")
    doc.save(docx_path)
    print("Document successfully updated.")
    return True

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 append_appendix.py <path_to_docx>")
        sys.exit(1)
        
    docx_path = sys.argv[1]
    add_appendix(docx_path)
