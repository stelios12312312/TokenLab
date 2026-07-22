# tests/test_docx_verification.py
# @planner:module = test_docx_verification
# @planner:story = US-Z1-M3-09
import os
import shutil
import pytest
import docx
from TokenLab.utils.docx_verifier import DocxVerifier
from TokenLab.utils.verifier import load_spec

def test_docx_verification_success():
    """
    US-Z1-M3-09: Verifies that the verifier runs successfully when comparing
    spec.yaml against docs_final/Z1_TOKEN_LIFECYCLE_V2.docx.
    """
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    spec_path = os.path.join(root, "projects", "z1", "spec.yaml")
    docx_path = os.path.join(root, "docs_final", "Z1_TOKEN_LIFECYCLE_V2.docx")
    
    spec = load_spec(spec_path)
    verifier = DocxVerifier(spec, docx_path)
    verifier.extract_parameters()
    success = verifier.verify_compliance()
    
    assert success, "DOCX compliance verification failed for the canonical Z1 files"
    # Ensure expected parameters are in docx_parameters
    assert "TAU_1" in verifier.docx_parameters
    assert "LM_RATE" in verifier.docx_parameters
    assert "settlement_cap_epoch" in verifier.docx_parameters

def test_docx_verification_drift_failure():
    """
    US-Z1-M3-09: Verifies that a drift in a parameter value in the DOCX
    correctly triggers a verification failure.
    """
    root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    spec_path = os.path.join(root, "projects", "z1", "spec.yaml")
    docx_path = os.path.join(root, "docs_final", "Z1_TOKEN_LIFECYCLE_V2.docx")
    
    mutated_docx_path = os.path.join(root, "tests", "temp_mutated_test.docx")
    try:
        shutil.copy(docx_path, mutated_docx_path)
        doc = docx.Document(mutated_docx_path)
        
        # Verify initial state of cell in Table 24 (Parameter Calibration Registry)
        # Table 24, Row 1 is TAU_1, cell 0 is "TAU_1", cell 2 is "0.20"
        assert doc.tables[24].rows[1].cells[0].text == "TAU_1"
        assert doc.tables[24].rows[1].cells[2].text.strip() == "0.20"
        
        # Mutate value to 0.99
        doc.tables[24].rows[1].cells[2].text = "0.99"
        doc.save(mutated_docx_path)
        
        spec = load_spec(spec_path)
        verifier = DocxVerifier(spec, mutated_docx_path)
        verifier.extract_parameters()
        success = verifier.verify_compliance()
        
        assert not success, "Verification should fail when a parameter value drifts"
        assert any("Mismatch for 'utility_fee_share'" in err for err in verifier.errors), "Errors should report the mismatch for utility_fee_share"
    finally:
        if os.path.exists(mutated_docx_path):
            os.remove(mutated_docx_path)
